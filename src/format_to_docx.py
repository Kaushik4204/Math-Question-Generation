import os
from docx import Document
from typing import List, Optional

def parse_tagged_question(raw_text: str) -> dict:
    # Placeholder — replace with your real parser or import it.
    raise NotImplementedError("Replace parse_tagged_question with your implementation.")

def save_questions_to_word(
    questions: List[dict],  # Now expects a list of parsed dicts, not raw strings
    filename: str,
    title: Optional[str] = "Generated Questions"
) -> None:
    if not filename.lower().endswith(".docx"):
        filename += ".docx"
    os.makedirs(os.path.dirname(filename) or ".", exist_ok=True)

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for idx, tags in enumerate(questions, start=1):
        doc.add_heading(f"Question {idx}", level=2)

        if tags.get("title"):
            doc.add_heading(tags["title"], level=3)

        if tags.get("description"):
            p = doc.add_paragraph(tags["description"])
            p.italic = True

        if tags.get("question"):
            doc.add_paragraph(tags["question"])

        if tags.get("instruction"):
            p = doc.add_paragraph(f"Instruction: {tags['instruction']}")
            p.italic = True

        # Metadata line
        metadata_parts = []
        for key in ["difficulty", "subject", "unit", "topic", "plusmarks"]:
            val = tags.get(key)
            if val:
                display_key = "Marks" if key == "plusmarks" else key.capitalize()
                metadata_parts.append(f"{display_key}: {val}")
        if metadata_parts:
            p = doc.add_paragraph(" | ".join(metadata_parts))
            p.italic = True

        # Remove duplicates preserving order
        seen = set()
        unique_options = []
        for opt in tags.get("options", []):
            if opt not in seen:
                unique_options.append(opt)
                seen.add(opt)

        correct_answer = tags.get("correct_answer")

        for opt in unique_options:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(opt)
            if opt == correct_answer:
                # Prepend checkmark and bold text properly
                run.text = ""
                run.add_text("✔️ ")
                run.add_text(opt)
                run.bold = True

        if tags.get("explanation"):
            p = doc.add_paragraph(f"Explanation: {tags['explanation']}")
            p.italic = True

        doc.add_paragraph("")  # blank line for spacing

    doc.save(filename)
    print(f"✅ Word file saved: {filename}")
